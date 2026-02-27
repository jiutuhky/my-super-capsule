#!/usr/bin/env python3
"""
insert_diagrams.py - 向专利申请 Word 文件中插入附图

将 PNG 附图插入到已填充文本的专利申请 .docx 文件的指定 Section 中。
由 docx-merger agent 分析内容后调用，传入明确的附图映射参数。

用法:
    # 插入摘要附图到 Section 1
    python3 insert_diagrams.py \
        --docx "06_final/patent_application.docx" \
        --section 1 \
        --figures "1:/path/to/fig1.png"

    # 插入说明书附图到 Section 4（按传入顺序插入）
    python3 insert_diagrams.py \
        --docx "06_final/patent_application.docx" \
        --section 4 \
        --figures "1:/path/to/fig1.png" "2:/path/to/fig2.png" "3:/path/to/fig3.png"

参数说明:
    --figures 格式为 "图号:文件路径"，可传入多个，按传入顺序插入。
    --section 1 表示摘要附图（仅插入图片，无标签）。
    --section 4 表示说明书附图（每张图前加居中图号标签）。
"""

import argparse
import os
import sys

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find_section_boundaries(doc):
    """找到每个 Section 的段落索引范围。"""
    paragraphs = doc.paragraphs
    boundaries = []
    for i, p in enumerate(paragraphs):
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            boundaries.append(i)

    sections = []
    start = 0
    for b in boundaries:
        sections.append((start, b, b))
        start = b + 1
    sections.append((start, len(paragraphs) - 1, None))
    return sections


def make_text_paragraph(text='', center=False, font_size_pt=14):
    """创建一个标准格式的文本段落（26pt 固定行距）。"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '520')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), '0')
    pPr.append(ind)

    if center:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(szCs)

        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)

    return p


def add_picture_paragraph(doc, image_path, max_width_cm=17):
    """
    创建一个包含图片的居中段落。
    使用单倍自动行距，避免固定行距裁剪图片。
    """
    p = doc.add_paragraph()
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)

    # 居中
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    # 单倍自动行距（避免固定行距裁剪图片）
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    # 无首行缩进
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), '0')
    pPr.append(ind)

    run = p.add_run()
    inline = run.add_picture(image_path, width=Cm(max_width_cm))

    # 按宽度上限等比缩放
    max_w = Cm(max_width_cm)
    if inline.width > max_w:
        ratio = max_w / inline.width
        inline.width = int(inline.width * ratio)
        inline.height = int(inline.height * ratio)

    return p


def parse_figure_arg(fig_str):
    """解析 '图号:路径' 格式的参数，返回 (fig_num_str, path)。"""
    parts = fig_str.split(':', 1)
    if len(parts) != 2:
        print(f'错误: 附图参数格式不正确: {fig_str}', file=sys.stderr)
        print('  正确格式: "图号:文件路径"，例如 "1:/path/to/fig1.png"', file=sys.stderr)
        sys.exit(1)
    fig_num, path = parts[0].strip(), parts[1].strip()
    if not os.path.exists(path):
        print(f'错误: 附图文件不存在: {path}', file=sys.stderr)
        sys.exit(1)
    return fig_num, path


# ---------------------------------------------------------------------------
# Section inserters
# ---------------------------------------------------------------------------

def insert_into_section_1(doc, figures):
    """
    Section 1: 摘要附图。
    仅插入图片（通常只有一张），无图号标签。
    插入位置：sectPr 段落之前。
    """
    sections = find_section_boundaries(doc)
    sec_start, sec_end, sectpr_idx = sections[1]
    sectpr_para = doc.paragraphs[sectpr_idx]._element

    for fig_num, fig_path in figures:
        pic_p = add_picture_paragraph(doc, fig_path)
        pic_elem = pic_p._element
        pic_elem.getparent().remove(pic_elem)
        sectpr_para.addprevious(pic_elem)

    print(f'  Section 1 (摘要附图): 已插入 {len(figures)} 张图片')


def insert_into_section_4(doc, figures):
    """
    Section 4: 说明书附图。
    每张图前加居中图号标签，按传入顺序插入。
    插入位置：body 级 sectPr 之前。
    """
    body = doc.element.body
    body_sectpr = body.find(qn('w:sectPr'))

    for fig_num, fig_path in figures:
        # 图号标签（居中）
        label_p = make_text_paragraph(f'图{fig_num}', center=True)
        body_sectpr.addprevious(label_p)

        # 空段落
        empty_p = make_text_paragraph('', center=False)
        body_sectpr.addprevious(empty_p)

        # 图片
        pic_p = add_picture_paragraph(doc, fig_path)
        pic_elem = pic_p._element
        pic_elem.getparent().remove(pic_elem)
        body_sectpr.addprevious(pic_elem)

        # 图后空段落
        empty_p2 = make_text_paragraph('', center=False)
        body_sectpr.addprevious(empty_p2)

    print(f'  Section 4 (说明书附图): 已插入 {len(figures)} 张图片')


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='向专利申请 Word 文件中插入附图'
    )
    parser.add_argument('--docx', required=True, help='已填充文本的 .docx 路径')
    parser.add_argument('--section', required=True, type=int, choices=[1, 4],
                        help='目标 Section: 1=摘要附图, 4=说明书附图')
    parser.add_argument('--figures', required=True, nargs='+',
                        help='附图列表，格式: "图号:文件路径"，可传入多个')

    args = parser.parse_args()

    if not os.path.exists(args.docx):
        print(f'错误: .docx 文件不存在: {args.docx}', file=sys.stderr)
        sys.exit(1)

    figures = [parse_figure_arg(f) for f in args.figures]

    print(f'正在加载: {args.docx}')
    doc = Document(args.docx)

    sections = find_section_boundaries(doc)
    if len(sections) != 5:
        print(f'错误: 文档应有 5 个 section，实际有 {len(sections)} 个',
              file=sys.stderr)
        sys.exit(1)

    print(f'正在插入附图到 Section {args.section}...')
    if args.section == 1:
        insert_into_section_1(doc, figures)
    elif args.section == 4:
        insert_into_section_4(doc, figures)

    doc.save(args.docx)
    print(f'已保存: {args.docx}')

    # 统计
    result_doc = Document(args.docx)
    total_images = len(result_doc.inline_shapes)
    print(f'文档中共有 {total_images} 张图片')


if __name__ == '__main__':
    main()
