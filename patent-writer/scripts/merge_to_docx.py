#!/usr/bin/env python3
"""
merge_to_docx.py - 基于 Word 模板生成专利申请文件（纯文本部分）

将 abstract.md, claims.md, description.md 填充到专利申请模板的
Section 0（说明书摘要）、Section 2（权利要求书）、Section 3（说明书）中。
Section 1（摘要附图）和 Section 4（说明书附图）留空，由后续步骤插入附图。

用法:
    python3 merge_to_docx.py \
        --template "skills/writing-patent/references/template.docx" \
        --abstract "04_content/abstract.md" \
        --claims "04_content/claims.md" \
        --description "04_content/description.md" \
        --output "06_final/patent_application.docx"
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find_section_boundaries(doc):
    """
    找到每个 Section 的段落索引范围。

    返回: [(start_idx, end_idx, sectPr_paragraph_idx_or_None), ...]
    前 N-1 个 section 的边界由段落内嵌 w:sectPr 定义，
    最后一个 section 的 sectPr 在 body 级别。
    """
    paragraphs = doc.paragraphs
    boundaries = []
    for i, p in enumerate(paragraphs):
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            boundaries.append(i)

    sections = []
    start = 0
    for b in boundaries:
        sections.append((start, b, b))
        start = b + 1
    sections.append((start, len(paragraphs) - 1, None))
    return sections


def make_paragraph_element(text='', bold=False, center=False,
                           first_line_indent=True, font_size_pt=14):
    """
    创建一个新的 w:p 元素，带标准专利格式。

    格式规范: 宋体/Times New Roman, 14pt, 行距 26pt 固定值,
    首行缩进 28pt (可选)。
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # 行距: 26pt 固定值 = 520 twips
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '520')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

    # 缩进
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), '560' if first_line_indent else '0')
    pPr.append(ind)

    # 对齐
    if center:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

    p.append(pPr)

    if text:
        _add_runs_to_paragraph(p, text, bold, font_size_pt)

    return p


def _add_runs_to_paragraph(p_elem, text, bold=False, font_size_pt=14):
    """给段落元素添加 run（处理 **粗体** 标记）。"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        is_bold = bold
        actual_text = part
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            actual_text = part[2:-2]

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(szCs)

        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)

        r.append(rPr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = actual_text
        r.append(t)

        p_elem.append(r)


def _collect_and_remove(doc, sec_start, sec_end, sectpr_idx):
    """收集并删除一个 section 内的非 sectPr 段落，返回 sectPr 段落元素。"""
    sectpr_para = doc.paragraphs[sectpr_idx]._element
    to_remove = []
    for i in range(sec_start, sec_end + 1):
        if i == sectpr_idx:
            continue
        to_remove.append(doc.paragraphs[i]._element)
    for elem in to_remove:
        elem.getparent().remove(elem)
    return sectpr_para


def _clear_last_section(doc, sec_start, sec_end):
    """清空最后一个 section 的段落（body 级 sectPr，不在段落内）。"""
    to_remove = []
    for i in range(sec_start, sec_end + 1):
        to_remove.append(doc.paragraphs[i]._element)
    for elem in to_remove:
        elem.getparent().remove(elem)


# ---------------------------------------------------------------------------
# Markdown Parsers
# ---------------------------------------------------------------------------

def parse_abstract(filepath):
    """
    解析 abstract.md。
    返回摘要正文（去掉标题行和"摘要附图"行）。
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    text_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('# '):
            continue
        if re.match(r'^摘要附图[：:]', stripped):
            continue
        text_lines.append(stripped)

    return ''.join(text_lines).strip()


def parse_claims(filepath):
    """
    解析 claims.md。
    返回权利要求列表 [(claim_number, claim_text), ...]。
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    body_lines = [l for l in lines if not l.strip().startswith('# ')]
    body = '\n'.join(body_lines).strip()

    claims = []
    current_num = None
    current_lines = []

    for line in body.split('\n'):
        m = re.match(r'^(\d+)\.\s', line)
        if m:
            if current_num is not None:
                claims.append((current_num, '\n'.join(current_lines).strip()))
            current_num = int(m.group(1))
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_num is not None:
        claims.append((current_num, '\n'.join(current_lines).strip()))

    return claims


def parse_description(filepath):
    """
    解析 description.md。
    返回 {'invention_name': str, 'sections': [{'title', 'level', 'paragraphs'}, ...]}
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    result = {'invention_name': '', 'sections': []}
    current_section = None
    in_invention_name = False
    invention_name_next = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('# ') and not stripped.startswith('## '):
            continue

        if stripped.startswith('## '):
            title = stripped[3:].strip()
            if title == '发明名称':
                in_invention_name = True
                invention_name_next = True
                continue
            else:
                in_invention_name = False
                if current_section:
                    result['sections'].append(current_section)
                current_section = {'title': title, 'level': 2, 'paragraphs': []}
                continue

        if stripped.startswith('### '):
            title = stripped[4:].strip()
            if current_section:
                result['sections'].append(current_section)
            current_section = {'title': title, 'level': 3, 'paragraphs': []}
            continue

        if in_invention_name and invention_name_next and stripped:
            result['invention_name'] = stripped
            invention_name_next = False
            continue

        if current_section is not None and stripped:
            current_section['paragraphs'].append(stripped)

    if current_section:
        result['sections'].append(current_section)

    return result


# ---------------------------------------------------------------------------
# Section Fillers
# ---------------------------------------------------------------------------

def fill_section_0_abstract(doc, sections_info, abstract_text):
    """Section 0: 说明书摘要"""
    sec_start, sec_end, sectpr_idx = sections_info[0]
    sectpr_para = _collect_and_remove(doc, sec_start, sec_end, sectpr_idx)
    p = make_paragraph_element(abstract_text, first_line_indent=True)
    sectpr_para.addprevious(p)


def fill_section_1_clear(doc):
    """Section 1: 摘要附图 — 清空模板占位内容，留空待后续插入。"""
    sections = find_section_boundaries(doc)
    sec_start, sec_end, sectpr_idx = sections[1]
    _collect_and_remove(doc, sec_start, sec_end, sectpr_idx)


def fill_section_2_claims(doc, claims):
    """Section 2: 权利要求书"""
    sections = find_section_boundaries(doc)
    sec_start, sec_end, sectpr_idx = sections[2]
    sectpr_para = _collect_and_remove(doc, sec_start, sec_end, sectpr_idx)

    for claim_num, claim_text in claims:
        paragraphs = claim_text.split('\n')
        merged = []
        current = []
        for line in paragraphs:
            if line.strip():
                current.append(line.strip())
            else:
                if current:
                    merged.append(' '.join(current))
                    current = []
        if current:
            merged.append(' '.join(current))

        for para_text in merged:
            p = make_paragraph_element(para_text, first_line_indent=False)
            sectpr_para.addprevious(p)

        empty_p = make_paragraph_element('', first_line_indent=False)
        sectpr_para.addprevious(empty_p)


def fill_section_3_description(doc, desc_data):
    """Section 3: 说明书"""
    sections = find_section_boundaries(doc)
    sec_start, sec_end, sectpr_idx = sections[3]
    sectpr_para = _collect_and_remove(doc, sec_start, sec_end, sectpr_idx)

    # 发明名称（居中加粗）
    name_p = make_paragraph_element(
        desc_data['invention_name'],
        bold=True, center=True, first_line_indent=False
    )
    sectpr_para.addprevious(name_p)

    # 各子节
    for section in desc_data['sections']:
        title_p = make_paragraph_element(
            section['title'], bold=True, first_line_indent=False
        )
        sectpr_para.addprevious(title_p)

        for para_text in section['paragraphs']:
            if re.match(r'^\*\*.*\*\*$', para_text):
                clean = para_text.strip('*').strip()
                content_p = make_paragraph_element(
                    clean, bold=True, first_line_indent=False
                )
            else:
                content_p = make_paragraph_element(
                    para_text, first_line_indent=True
                )
            sectpr_para.addprevious(content_p)


def fill_section_4_clear(doc):
    """Section 4: 说明书附图 — 清空模板占位内容，留空待后续插入。"""
    sections = find_section_boundaries(doc)
    sec_start, sec_end, sectpr_idx = sections[4]
    _clear_last_section(doc, sec_start, sec_end)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='基于 Word 模板生成专利申请文件（纯文本部分）'
    )
    parser.add_argument('--template', required=True, help='模板 .docx 路径')
    parser.add_argument('--abstract', required=True, help='abstract.md 路径')
    parser.add_argument('--claims', required=True, help='claims.md 路径')
    parser.add_argument('--description', required=True, help='description.md 路径')
    parser.add_argument('--output', required=True, help='输出 .docx 路径')

    args = parser.parse_args()

    for path, name in [
        (args.template, '模板文件'),
        (args.abstract, 'abstract.md'),
        (args.claims, 'claims.md'),
        (args.description, 'description.md'),
    ]:
        if not os.path.exists(path):
            print(f'错误: {name}不存在: {path}', file=sys.stderr)
            sys.exit(1)

    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)

    print('正在解析输入文件...')
    abstract_text = parse_abstract(args.abstract)
    claims = parse_claims(args.claims)
    desc_data = parse_description(args.description)

    print(f'  摘要: {len(abstract_text)} 字')
    print(f'  权利要求: {len(claims)} 条')
    print(f'  说明书章节: {len(desc_data["sections"])} 节')

    print('正在加载模板...')
    doc = Document(args.template)

    sections = find_section_boundaries(doc)
    if len(sections) != 5:
        print(f'错误: 模板应有 5 个 section，实际有 {len(sections)} 个',
              file=sys.stderr)
        sys.exit(1)

    print('正在填充 Section 0: 说明书摘要...')
    fill_section_0_abstract(doc, sections, abstract_text)

    print('正在清空 Section 1: 摘要附图（待后续插入）...')
    fill_section_1_clear(doc)

    print('正在填充 Section 2: 权利要求书...')
    fill_section_2_claims(doc, claims)

    print('正在填充 Section 3: 说明书...')
    fill_section_3_description(doc, desc_data)

    print('正在清空 Section 4: 说明书附图（待后续插入）...')
    fill_section_4_clear(doc)

    print(f'正在保存: {args.output}')
    doc.save(args.output)
    print('完成！文本内容已填充，附图 Section 留空待后续步骤插入。')

    result_doc = Document(args.output)
    print(f'\n输出文件统计:')
    print(f'  段落数: {len(result_doc.paragraphs)}')
    print(f'  Section 数: {len(result_doc.sections)}')
    for i, sec in enumerate(result_doc.sections):
        header_text = ''.join(p.text for p in sec.header.paragraphs)
        print(f'  Section {i} 页眉: {header_text}')


if __name__ == '__main__':
    main()
